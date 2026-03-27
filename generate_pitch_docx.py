
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

def convert_to_png(img_path, tmp_dir):
    """Converts an image to a temporary PNG file to ensure docx compatibility."""
    if not os.path.exists(img_path):
        return None
    try:
        img = Image.open(img_path)
        filename = os.path.basename(img_path).split('.')[0] + "_converted.png"
        tmp_path = os.path.join(tmp_dir, filename)
        img.save(tmp_path, "PNG")
        return tmp_path
    except Exception as e:
        print(f"Error converting {img_path}: {e}")
        return None

def create_pitch_doc():
    doc = Document()
    
    # Path setup
    artifact_dir = r"C:\Users\kushw\.gemini\antigravity\brain\fc57b811-99c9-41eb-bc18-991523fcea64"
    tmp_dir = r"c:\Users\kushw\OneDrive\Desktop\antigravity\booth-iq-final\latest-25th-march-booth-iq-repo\tmp"
    if not os.path.exists(tmp_dir):
        os.makedirs(tmp_dir)
        
    img_ps_orig = os.path.join(artifact_dir, "booth_iq_problem_statement_visual_1774626333959.png")
    img_stack_orig = os.path.join(artifact_dir, "booth_iq_details_stack_light_2d_1774625756993.png")
    img_pipeline_orig = os.path.join(artifact_dir, "booth_iq_ai_pipeline_light_2d_1774625820265.png")
    
    # Pre-convert images
    img_ps = convert_to_png(img_ps_orig, tmp_dir)
    img_stack = convert_to_png(img_stack_orig, tmp_dir)
    img_pipeline = convert_to_png(img_pipeline_orig, tmp_dir)
    
    # Title Page
    title = doc.add_heading('BoothIQ: Technical Pitch & Operational Strategy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Strategic Presentation for Judges & Governance Experts').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Problem Statement Visual to Title/Intro
    if img_ps:
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_ps, width=Inches(5.5))
        doc.add_paragraph("Figure 1: The Digital Chasm - From Static Lists to Living Intelligence").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Section 1: The Pitch Script
    doc.add_heading('🎙️ 1. The Official Pitch Script', level=1)
    
    sections = [
        ("1. The Hook", "Ladies and Gentlemen, we aren't just missing data; we’re missing Intelligence at the Edge."),
        ("1.1 The Problem Statement (The 'PS')", "How do we transform static, siloed voter lists into a living knowledge system? BoothIQ bridges the 'Digital Chasm' by moving from rows in a spreadsheet to real-time execution oversight."),
        ("2. The Solution", "Introducing BoothIQ—a Hybrid Intelligence platform designed to operationalize governance at the last mile."),
        ("3. The Demo: Citizen Reporter", "Utilizing Sarvam AI and GPT-4o-mini for hyper-localized Speech-to-Text and classification."),
        ("4. The Demo: The Command Matrix", "Real-time sentiment volatility, Knowledge Graphs, and Social Fabric mapping."),
        ("5. The Technical Moat", "Hybrid Data Resilience (Supabase + MongoDB), AI-First Middleware, and the Social Fabric Algorithm."),
        ("6. Impact & Scaling", "Moving from Passive Representation to Predictive Governance.")
    ]
    
    for head, text in sections:
        h = doc.add_heading(head, level=2)
        p = doc.add_paragraph(text)
        p.italic = True
    
    doc.add_page_break()
    
    # Section 2: Technical Stack
    doc.add_heading('🛠️ 2. Technical Stack', level=1)
    
    if img_stack:
        doc.add_picture(img_stack, width=Inches(6.0))
        doc.add_paragraph("Figure 2: Comprehensive Technology Stack Architecture").alignment = WD_ALIGN_PARAGRAPH.CENTER

    tech_details = [
        ("Frontend", "React 19, Vite, Framer Motion, Tailwind CSS"),
        ("Backend", "FastAPI (Python 3.11), Pydantic V2, WebSockets"),
        ("Data Layer", "Supabase (Relational), MongoDB Atlas (Document/Graph)"),
        ("Intelligence", "OpenAI (GPT-4o-mini), Sarvam AI (STT/TTS)"),
        ("Communication", "Twilio (SMS/WhatsApp), Resend (Email)")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Used'
    for comp, tech in tech_details:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    doc.add_page_break()
    
    # Section 3: AI Pipeline
    doc.add_heading('🧠 3. AI & Machine Learning Pipeline', level=1)
    
    if img_pipeline:
        doc.add_picture(img_pipeline, width=Inches(6.0))
        doc.add_paragraph("Figure 3: Multi-Modal AI Pipeline Flow").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("- Grievance Classification: GPT-4o-mini for real-time prioritization.")
    doc.add_paragraph("- Voice Interface: Sarvam-m for localized Indian language processing.")
    doc.add_paragraph("- Social Fabric mapping: Knowledge Graph linkage based on household clustering.")
    
    doc.add_page_break()
    
    # Section 4: Knowledge Graph & Social Fabric: Deep Dive
    doc.add_heading('🧩 4. Knowledge Graph & Social Fabric: Deep Dive', level=1)
    
    kg_qa = [
        ("What is the BoothIQ Knowledge Graph?", 
         "It is a non-linear data structure representing voters as nodes and their relationships (family clusters, geographic proximity, and engagement history) as edges. While traditional databases store 'isolated data', the Knowledge Graph stores 'interconnected reality'."),
        
        ("Why use a Knowledge Graph instead of a standard SQL database?", 
         "Standard databases (SQL) are designed for tables and rows. Querying relationships (e.g., 'Who are all the influencers in this building?') requires complex, slow joins. The Knowledge Graph (powered by MongoDB Document/Graph models) allows us to traverse relationships at millisecond speed, crucial for real-time command centers."),
        
        ("How does the 'Social Fabric' mapping actually work?", 
         "We use a proprietary algorithm that links voters based on: 1. Household ID (Family unit), 2. Geographic Clustering (Building/Street proximity), and 3. Historical Interaction (Call logs and grievance patterns). This creates a 'Map of Influence'."),
        
        ("What is the primary reason for using it in governance?", 
         "It transitions us from 'Passive Reporting' to 'Predictive Action'. If we see a grievance from a community 'influencer' at Booth 17, we prioritize it because satisfying that one node positively impacts 50 other connected nodes in the graph. It's about 'Strategic Resource Optimization'."),
        
        ("How does the AI interact with the Graph?", 
         "Our LLM (ESarthi) queries the graph to provide tactical advice. Instead of just saying 'Grievance received', it says: 'This issue is from a high-influence household; resolution here will likely trigger a +5% sentiment shift in the ward.'")
    ]
    
    for q, a in kg_qa:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(f"Ans: {a}")
        doc.add_paragraph() # Spacer
        
    doc.add_page_break()
    
    # Section 5: Judges' FAQ
    doc.add_heading('💡 5. Judges\' FAQ: Technical Defenses', level=1)
    faqs = [
        ("How do you handle fake reports?", "AI Vision analyzes photo attachments; Reporter Trust Scores based on resolution history."),
        ("Why MongoDB AND Supabase?", "Supabase for relational data; MongoDB for high-velocity graphs."),
        ("Is it privacy-compliant?", "JWT-based RBAC ensures PII is only visible to authorized personnel."),
        ("How do you handle low connectivity?", "Queued Dispatcher pattern; asynchronous voice transcription.")
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(f"Q: {q}").bold = True
        doc.add_paragraph(f"Ans: {a}")
        
    doc.add_page_break()
    
    # Section 6: Loopholes
    doc.add_heading('🛡️ 6. Security & Mitigation', level=1)
    loopholes = [
        ("Zero-Password Auth", "Implement OAuth 2.0 or OTP-based verification."),
        ("Universal IDOR", "Add owner-validation decorators to endpoints."),
        ("Hardcoded Secrets", "Rotate secrets and enforce environment-only config."),
        ("Demo Bypass", "Strictly reject unauthenticated requests in production."),
        ("Regex AI Manipulation", "Use LLM-based verification instead of simple RegEx.")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Loophole'
    hdr_cells[1].text = 'Tackle/Fix'
    for loop, fix in loopholes:
        row_cells = table.add_row().cells
        row_cells[0].text = loop
        row_cells[1].text = fix

    doc.save('BoothIQ_Pitch_Deck_Final.docx')
    print("Document saved with All Context as BoothIQ_Pitch_Deck_Final.docx")

if __name__ == "__main__":
    create_pitch_doc()
